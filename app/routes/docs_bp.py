import os
import shutil
import time
import zipfile
from flask import Blueprint, redirect, send_file, url_for, after_this_request
from ..data_base import *
from docx import Document
from docx.table import Table
from re import search
from datetime import date

docs_bp = Blueprint('Blueprint_docs', __name__)

def replace_text(doc, repor):
   for paragrafo in doc.paragraphs:
      for antigo, novo in repor.items():
         if antigo in paragrafo.text:
               paragrafo.text = paragrafo.text.replace(antigo, novo)

   for tabela in doc.tables:
      for linha in tabela.rows:
         for celula in linha.cells:
               for antigo, novo in repor.items():
                  if antigo in celula.text:
                     celula.text = celula.text.replace(antigo, novo)

@docs_bp.route('/Baixar/<int:estagio_id>', methods=["POST", "GET"])
def baixar(estagio_id):
   estagio = session.query(Estagios).filter_by(id=estagio_id).first()
   aluno = session.query(Alunos).filter_by(id=estagio.alunoId).first()
   turma = session.query(Turmas).filter_by(id=aluno.turmaId).first()
   entidade = session.query(Entidade).filter_by(id=estagio.entidadeId).first()
   numero = search(r'\d+', turma.descricao).group()

   # Caminho onde o arquivo ZIP será salvo temporariamente
   output_folder = os.path.join(os.getcwd(), 'app', 'temp')

   # Verifica se a pasta temporária existe, caso contrário, cria-a
   if not os.path.exists(output_folder):
      os.makedirs(output_folder)

   repor = {
      '[nome_aluno]': aluno.nome,
      '[turma]': numero,
      '[nr_aluno]': aluno.nr,
      '[orientador]': estagio.orientador,
      '[tutor]': estagio.tutor,
      '[Nome da Entidade/Empresa]': entidade.nome,
      '[Número de Contribuinte Entidade]': entidade.nif,
      '[Morada Entidade]': entidade.morada,
      '[Nome Completo]': entidade.pessoa_responsavel,
      '[Cargo/Função]': entidade.cargo_pessoa_responsavel,
      '[Nome Completo do Aluno]': aluno.nome,
      '[Número do CC]': aluno.cartao_cidadao,
      '[dd/mm/aaaa]': aluno.validade_cc.strftime("%d/%m/%Y"),
      '[Número de Contribuinte Aluno]': aluno.nif,
      '[Morada Aluno]': aluno.morada,
      '[nome]': aluno.nome,
      '[data]': str(date.today().strftime("%d-%m-%Y"))
   }

   # Cria o arquivo ZIP
   zip_file_path = os.path.join(output_folder, f'documents - {aluno.nome}.zip')
   with zipfile.ZipFile(zip_file_path, 'w') as zip_file:
      doc1 = Document('app/docs/Folha de Presencas e Registo de Atividades Semanal.docx')
      replace_text(doc1, repor)
      doc1.save(os.path.join(output_folder, f'Folha de Presencas e Registo de Atividades Semanal - {aluno.nome}.docx'))
      zip_file.write(os.path.join(output_folder, f'Folha de Presencas e Registo de Atividades Semanal - {aluno.nome}.docx'), f'Folha de Presencas e Registo de Atividades Semanal - {aluno.nome}.docx')

      doc2 = Document('app/docs/Plano de FCT.docx')
      replace_text(doc2, repor)
      doc2.save(os.path.join(output_folder, f'Plano de FCT - {aluno.nome}.docx'))
      zip_file.write(os.path.join(output_folder, f'Plano de FCT - {aluno.nome}.docx'), f'Plano de FCT - {aluno.nome}.docx')

      doc3 = Document('app/docs/Grelha de Avaliacao da FCT.docx')
      replace_text(doc3, repor)
      doc3.save(os.path.join(output_folder, f'Grelha de Avaliacao da FCT - {aluno.nome}.docx'))
      zip_file.write(os.path.join(output_folder, f'Grelha de Avaliacao da FCT - {aluno.nome}.docx'), f'Grelha de Avaliacao da FCT - {aluno.nome}.docx')

      doc4 = Document('app/docs/Protocolo de FCT.docx')
      replace_text(doc4, repor)
      doc4.save(os.path.join(output_folder, f'Protocolo de FCT - {aluno.nome}.docx'))
      zip_file.write(os.path.join(output_folder, f'Protocolo de FCT - {aluno.nome}.docx'), f'Protocolo de FCT - {aluno.nome}.docx')

      doc5 = Document('app/docs/Plano Individual de FCT.docx')
      replace_text(doc5, repor)
      doc5.save(os.path.join(output_folder, f'Plano Individual de FCT - {aluno.nome}.docx'))
      zip_file.write(os.path.join(output_folder, f'Plano Individual de FCT - {aluno.nome}.docx'), f'Plano Individual de FCT - {aluno.nome}.docx')

   # Verifique se o arquivo ZIP foi criado corretamente
   if not os.path.exists(zip_file_path):
      return "Erro: Falha ao criar o arquivo ZIP"

   @after_this_request
   def remove_temp_folder(response):
      # Remove o arquivo individual
      os.remove(os.path.join(output_folder, f'Folha de Presencas e Registo de Atividades Semanal - {aluno.nome}.docx'))
      os.remove(os.path.join(output_folder, f'Plano de FCT - {aluno.nome}.docx'))
      os.remove(os.path.join(output_folder, f'Grelha de Avaliacao da FCT - {aluno.nome}.docx'))
      os.remove(os.path.join(output_folder, f'Protocolo de FCT - {aluno.nome}.docx'))
      os.remove(os.path.join(output_folder, f'Plano Individual de FCT - {aluno.nome}.docx'))
      
      return response

   # Envia o arquivo ZIP como resposta para o navegador
   return send_file(zip_file_path, as_attachment=True)